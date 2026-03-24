#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件处理器增强模块 - 集成OCR功能
"""
import os
import re
from pathlib import Path
from collections import Counter
from datetime import datetime
# OCR相关库
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("警告: OCR功能不可用，请安装 pytesseract, Pillow, pdf2image")

# 导入增强模块
try:
    from image_preprocessor import image_preprocessor
    from id_card_ocr import id_card_ocr
    ENHANCED_OCR = True
except ImportError:
    ENHANCED_OCR = False
    print("警告: 增强OCR模块不可用")

# Excel文件支持
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("警告: Excel文件解析不可用，请安装 openpyxl")

# Word文档支持
try:
    from docx import Document
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    print("警告: Word文档解析不可用，请安装 python-docx")

class FileProcessor:
    def process_file(self, file_path):
        """处理文件，返回处理结果 - 增强版，提升性能和准确性"""
        result = {
            'success': False,
            'content': '',
            'preview': '',
            'keywords': [],
            'worksheets_data': None,
            'metadata': {},
            'classification': {},  # 新增：分类信息
            'quality_score': 0.0,  # 新增：质量评分
        }
        
        try:
            if not os.path.exists(file_path):
                result['error'] = '文件不存在'
                return result
            
            # 第一步：提取文件元数据（快速，总是成功）
            metadata = self.extract_file_metadata(file_path)
            result['metadata'] = metadata
            
            # 提取分类信息
            classification = self.get_file_classification(file_path)
            result['classification'] = classification
            
            # 预先基于文件名/路径提取元数据关键词（最高优先级）
            metadata_keywords = self._extract_metadata_keywords(metadata)
            
            # 第二步：提取文本内容（可能需要OCR，较慢）
            content = self._extract_text_with_fallback(file_path)
            
            # 第三步：处理提取的内容
            if content and len(content.strip()) > 10:
                # 过滤水印和无意义内容
                filtered_content = self._filter_watermark(content)
                
                if filtered_content and len(filtered_content.strip()) > 10:
                    result['success'] = True
                    result['content'] = filtered_content
                    result['preview'] = self._generate_preview(filtered_content)
                    
                    # 根据优先级合并关键词：先文件名/路径，再正文内容
                    content_keywords = self.extract_keywords(filtered_content, max_keywords=50)
                    result['keywords'] = self._merge_keywords(metadata_keywords, content_keywords, max_total=30)
                    
                    # 计算质量评分
                    result['quality_score'] = self._calculate_quality_score(filtered_content, result['keywords'])
                else:
                    # 内容被过滤后为空，使用元数据
                    result['success'] = True
                    metadata_content = self._build_metadata_content(metadata)
                    result['content'] = "【内容提取失败，仅保存元数据】\n" + metadata_content
                    result['preview'] = result['content'][:500]
                    result['keywords'] = metadata_keywords
                    result['quality_score'] = 0.3  # 低质量评分
            else:
                # 无内容，使用元数据
                result['success'] = True
                metadata_content = self._build_metadata_content(metadata)
                result['content'] = "【内容提取失败，仅保存元数据】\n" + metadata_content
                result['preview'] = result['content'][:500]
                result['keywords'] = metadata_keywords
                result['quality_score'] = 0.3  # 低质量评分
                
        except Exception as e:
            print(f"处理文件失败: {e}")
            # 即使失败，也尝试提取元数据（增强的错误处理）
            try:
                metadata = self.extract_file_metadata(file_path)
                classification = self.get_file_classification(file_path)
                result['success'] = True
                metadata_content = self._build_metadata_content(metadata)
                result['content'] = "【内容提取失败，仅保存元数据】\n" + metadata_content
                result['preview'] = result['content'][:500]
                result['keywords'] = self._extract_metadata_keywords(metadata)
                result['metadata'] = metadata
                result['classification'] = classification
                result['quality_score'] = 0.2
                result['error'] = str(e)
            except Exception as e2:
                result['success'] = False
                result['error'] = f"全部失败: {str(e)}, {str(e2)}"
            
        return result
    
    def _extract_text_with_fallback(self, file_path):
        """增强的文本提取，支持多种回退策略"""
        try:
            # 尝试标准提取
            content = self.extract_text(file_path)
            
            # 如果内容过少且是图片文件，尝试使用增强的OCR
            if len(content.strip()) < 20:
                file_ext = Path(file_path).suffix.lower()
                if file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
                    content = self._enhanced_ocr_extraction(file_path)
            
            return content
            
        except Exception as e:
            print(f"文本提取失败: {e}")
            return ""
    
    def _enhanced_ocr_extraction(self, file_path):
        """增强的OCR提取 - 处理低质量图片"""
        if not OCR_AVAILABLE:
            return ""
        
        try:
            # 尝试多种图片预处理策略
            strategies = [
                'original',      # 原始图片
                'grayscale',     # 灰度化
                'binary',        # 二值化
                'denoise',       # 去噪
                'enhance',       # 增强对比度
            ]
            
            best_result = ""
            max_length = 0
            
            for strategy in strategies:
                try:
                    # 根据策略预处理图片
                    if strategy == 'original':
                        image = Image.open(file_path)
                    elif ENHANCED_OCR:
                        # 使用增强的预处理器
                        image = image_preprocessor.preprocess(file_path, strategy=strategy)
                    else:
                        # 简单预处理
                        image = self._simple_preprocess(file_path, strategy)
                    
                    # OCR识别
                    text = pytesseract.image_to_string(image, lang='chi_sim+eng', config='--psm 3 --oem 1')
                    
                    # 选择识别结果最多的
                    if len(text) > max_length:
                        best_result = text
                        max_length = len(text)
                    
                    # 如果已经获得足够的内容，不再尝试其他策略
                    if len(text.strip()) > 100:
                        break
                        
                except Exception as e:
                    print(f"OCR策略 {strategy} 失败: {e}")
                    continue
            
            return best_result.strip()
            
        except Exception as e:
            print(f"增强的OCR提取失败: {e}")
            return ""
    
    def _simple_preprocess(self, file_path, strategy):
        """简单的图片预处理（当增强模块不可用时）"""
        try:
            from PIL import ImageEnhance, ImageFilter
            image = Image.open(file_path)
            
            if strategy == 'grayscale':
                image = image.convert('L')
            elif strategy == 'binary':
                image = image.convert('L')
                # 简单二值化
                threshold = 128
                image = image.point(lambda x: 0 if x < threshold else 255, '1')
            elif strategy == 'denoise':
                image = image.filter(ImageFilter.MedianFilter())
            elif strategy == 'enhance':
                image = image.convert('L')
                enhancer = ImageEnhance.Contrast(image)
                image = enhancer.enhance(2.0)
            
            return image
            
        except Exception as e:
            print(f"图片预处理失败: {e}")
            return Image.open(file_path)
    
    def _generate_preview(self, content, max_length=500):
        """生成内容预览 - 智能截取"""
        if len(content) <= max_length:
            return content
        
        # 尝试在句子边界截断
        preview = content[:max_length]
        
        # 查找最后一个句子结束符
        last_period = max(
            preview.rfind('。'),
            preview.rfind('!'),
            preview.rfind('?'),
            preview.rfind('\n')
        )
        
        if last_period > max_length * 0.7:  # 如果截断点在后70%位置，使用它
            preview = content[:last_period + 1]
        
        return preview + '...'
    
    def _calculate_quality_score(self, content, keywords):
        """计算内容质量评分 (0-1)"""
        score = 0.0
        
        # 1. 内容长度评分（0-0.3）
        content_length = len(content.strip())
        if content_length >= 500:
            score += 0.3
        elif content_length >= 200:
            score += 0.2
        elif content_length >= 50:
            score += 0.1
        
        # 2. 关键词数量评分（0-0.2）
        keyword_count = len(keywords)
        if keyword_count >= 10:
            score += 0.2
        elif keyword_count >= 5:
            score += 0.1
        
        # 3. 中文占比评分（0-0.2）
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', content))
        chinese_ratio = chinese_chars / len(content) if len(content) > 0 else 0
        if chinese_ratio > 0.3:
            score += 0.2
        elif chinese_ratio > 0.1:
            score += 0.1
        
        # 4. 结构化程度评分（0-0.15）
        # 检查是否有段落结构
        lines = content.split('\n')
        non_empty_lines = [line for line in lines if line.strip()]
        if len(non_empty_lines) >= 5:
            score += 0.15
        elif len(non_empty_lines) >= 3:
            score += 0.1
        
        # 5. 特殊字符比例评分（0-0.15）
        # 太多特殊字符可能表示OCR质量差
        special_chars = len(re.findall(r'[^\w\s\u4e00-\u9fff]', content))
        special_ratio = special_chars / len(content) if len(content) > 0 else 0
        if special_ratio < 0.1:
            score += 0.15
        elif special_ratio < 0.2:
            score += 0.1
        
        return min(score, 1.0)  # 确保不超过1.0
    
    def get_file_type(self, file_path):
        """获取文件类型"""
        if not os.path.exists(file_path):
            return 'unknown'
        
        ext = Path(file_path).suffix.lower()
        
        # 简单的文件类型判断
        if ext in ['.pdf', '.doc', '.docx', '.txt', '.md']:
            return 'document'
        elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            return 'image'
        elif ext in ['.py', '.js', '.java', '.cpp', '.c', '.go']:
            return 'code'
        else:
            return 'other'
    
    def extract_text(self, file_path):
        """提取文件文本内容，支持OCR"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            # 纯文本文件
            if file_ext in ['.txt', '.md', '.log', '.csv']:
                return self._extract_plain_text(file_path)
            
            # 代码文件
            elif file_ext in ['.py', '.js', '.java', '.cpp', '.c', '.go', '.sh', '.html', '.css', '.json', '.xml', '.yaml', '.yml']:
                return self._extract_plain_text(file_path)
            
            # 图片文件 - OCR
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.gif']:
                return self._extract_image_text(file_path)
            
            # PDF文件 - OCR
            elif file_ext == '.pdf':
                return self._extract_pdf_text(file_path)
            
            # Excel文件
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_excel_text(file_path)
            
            # Word文档
            elif file_ext in ['.docx', '.doc']:
                return self._extract_word_text(file_path)
            
            # 其他Office文档（简单处理）
            elif file_ext in ['.ppt', '.pptx']:
                return f"文档类型: {file_ext}"
            
        except Exception as e:
            print(f"提取文件内容失败: {e}")
        return ""
    
    def _extract_plain_text(self, file_path):
        """提取纯文本内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
        except Exception as e:
            print(f"读取纯文本失败: {e}")
        return ""  # 失败时返回空
    
    def _extract_image_text(self, file_path):
        """使用OCR提取图片中的文字 - 增强版"""
        if not OCR_AVAILABLE:
            return ""  # 返回空，让上层处理
        
        try:
            # 判断是否为身份证类图片（根据文件路径或名称）
            is_id_card = self._is_id_card_image(file_path)
            
            if is_id_card and ENHANCED_OCR:
                # 使用专用身份证识别
                result = id_card_ocr.recognize(file_path)
                if result['success']:
                    # 返回处理后的文本
                    return result['processed_text']
            
            # 使用增强的OCR识别
            if ENHANCED_OCR:
                # 图片预处理
                image = image_preprocessor.preprocess(file_path)
            else:
                image = Image.open(file_path)
            
            # 使用多种PSM模式尝试识别
            psm_modes = [3, 6, 11]  # 3=自动, 6=单列文本, 11=稀疏文本
            best_text = ''
            max_length = 0
            
            for psm in psm_modes:
                try:
                    custom_config = f'--psm {psm} --oem 1'
                    text = pytesseract.image_to_string(image, lang='chi_sim+eng', config=custom_config)
                    # 选择识别内容最多的结果
                    if len(text) > max_length:
                        best_text = text
                        max_length = len(text)
                except:
                    continue
            
            return best_text.strip() if best_text else ""  # 无内容时返回空
            
        except Exception as e:
            print(f"OCR提取图片文字失败: {e}")
            return ""  # 失败时返回空
    
    def _is_id_card_image(self, file_path):
        """判断是否为身份证图片"""
        # 根据路径和文件名判断
        path_lower = file_path.lower()
        keywords = ['身份证', 'idcard', 'id_card', '证件', 'card']
        return any(keyword in path_lower for keyword in keywords)
    
    def _extract_excel_text(self, file_path):
        """提取Excel文件内容"""
        if not EXCEL_AVAILABLE:
            return ""  # 返回空，让上层处理
        
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = f"\n=== 工作表: {sheet_name} ===\n"
                
                rows_data = []
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空行
                    if any(cell is not None for cell in row):
                        # 转换为字符串并过滤None
                        row_str = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                        if row_str.strip():
                            rows_data.append(row_str)
                
                if rows_data:
                    sheet_text += '\n'.join(rows_data)
                    text_parts.append(sheet_text)
            
            workbook.close()
            return '\n'.join(text_parts) if text_parts else ""  # 无内容时返回空
            
        except Exception as e:
            print(f"提取Excel内容失败: {e}")
            return ""  # 失败时返回空
    
    def _extract_word_text(self, file_path):
        """提取Word文档内容"""
        if not WORD_AVAILABLE:
            return ""  # 返回空，让上层处理
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # 提取正文段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                table_text = "\n=== 表格 ===\n"
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text += row_text + '\n'
                if len(table_text) > 20:  # 有内容才添加
                    text_parts.append(table_text)
            
            return '\n'.join(text_parts) if text_parts else ""  # 无内容时返回空
            
        except Exception as e:
            print(f"提取Word内容失败: {e}")
            return ""  # 失败时返回空
    
    def _extract_pdf_text(self, file_path):
        """使用OCR提取PDF中的文字"""
        if not OCR_AVAILABLE:
            return ""  # 返回空，让上层处理
        
        try:
            # 将PDF转换为图片
            images = convert_from_path(file_path, dpi=200, first_page=1, last_page=5)  # 只处理前5页
            
            text_parts = []
            for i, image in enumerate(images):
                try:
                    text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                    if text.strip():
                        text_parts.append(f"--- 第{i+1}页 ---\n{text.strip()}")
                except Exception as e:
                    print(f"处理PDF第{i+1}页失败: {e}")
            
            return "\n\n".join(text_parts) if text_parts else ""  # 无内容时返回空
        except Exception as e:
            print(f"OCR提取PDF文字失败: {e}")
            return ""  # 失败时返回空
    
    def extract_keywords(self, text, max_keywords=20):
        """从文本中提取关键词（增强版） - 支持多层级分类和中文语义理解"""
        if not text or len(text.strip()) == 0:
            return []
            
        try:
            # 第一步：智能分词和预处理
            words = self._intelligent_tokenize(text)
                
            # 第二步：应用增强的停用词过滤
            filtered_words = self._filter_stopwords(words)
                
            if not filtered_words:
                return []
                
            # 第三步：识别专业术语和重要词汇
            enhanced_words = self._enhance_with_terminology(filtered_words, text)
                
            # 第四步：词频统计和权重计算
            word_scores = self._calculate_word_scores(enhanced_words, text)
                
            # 第五步：返回按权重排序的关键词
            sorted_keywords = sorted(word_scores.items(), key=lambda x: x[1], reverse=True)
            return [word for word, score in sorted_keywords[:max_keywords]]
                
        except Exception as e:
            print(f"提取关键词失败: {e}")
            return []
        
    def _intelligent_tokenize(self, text):
        """智能分词 - 增强中文语义理解"""
        words = []
            
        # 1. 提取中文词汇（2-4字词组，常见的中文词汇长度）
        chinese_pattern = r'[\u4e00-\u9fff]{2,4}'
        chinese_words = re.findall(chinese_pattern, text)
        words.extend(chinese_words)
            
        # 2. 提取英文单词（保留大小写信息用于专业术语识别）
        english_pattern = r'\b[A-Za-z][A-Za-z0-9]*\b'
        english_words = re.findall(english_pattern, text)
        words.extend(english_words)
            
        # 3. 提取数字+单位组合（如：2023年、100万、3个月）
        number_unit_pattern = r'\d+[年月日元万亿米千克吨个台套次]'
        number_units = re.findall(number_unit_pattern, text)
        words.extend(number_units)
            
        # 4. 提取专业编号（如：NO.123、编号-456、A001）
        code_pattern = r'[A-Z]{1,3}[0-9]{2,6}|NO\.[0-9]+|编号[0-9]+'
        codes = re.findall(code_pattern, text, re.IGNORECASE)
        words.extend(codes)
            
        return words
        
    def _filter_stopwords(self, words):
        """增强的停用词过滤 - 更精确的无意义词汇过滤"""
        # 扩展的停用词库
        stop_words = {
            # 英文常见虚词（扩展）
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'be',
            'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will',
            'would', 'should', 'could', 'may', 'might', 'must', 'can', 'this',
            'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they',
            'am', 'pm', 'rr', 'oe', 'rt', 'rn', 'nn', 'mm', 'll',
                
            # 常见中文虚词（扩展）
            '的', '了', '在', '是', '我', '有', '和', '人', '这', '个',
            '上', '中', '大', '为', '来', '不', '到', '说', '以', '及',
            '与', '或', '但', '而', '之', '其', '所', '于', '由', '对',
            '把', '从', '向', '往', '给', '让', '被', '将', '要', '会',
            '能', '可', '得', '着', '过', '去', '也', '都', '就', '只',
            '很', '太', '更', '最', '非常', '十分', '比较', '相当',
                
            # OCR常见错误识别词
            'rr', 'oe', 'rt', 'rn', 'nn', 'mm', 'll', 'ii', 'oo', 'vv',
            'rrr', 'ooo', 'nnn', 'mmm', 'lll',
                
            # 常见格式性/无意义词（扩展）
            '文件', '内容', '表格', '附件', '资料', '材料', '清单', '流程', '说明',
            '第', '页', '页面', '首页', '尾页', '封面', '目录', '索引',
            '公司', '部门', '日期', '编号', '序号', '标题', '备注', '摘要',
            '文档', '报告', '通知', '信息', '数据', '记录', '详情', '情况',
            '相关', '有关', '关于', '根据', '按照', '依据', '进行', '实施',
            '完成', '开展', '推进', '落实', '执行', '办理', '处理', '管理',
                
            # 页码和格式标记
            '第1页', '第2页', '第3页', '第4页', '第5页', '第6页', '第7页',
            '第8页', '第9页', '第10页', '页码', '共页', '总页数',
        }
            
        filtered = []
        for word in words:
            # 长度过滤
            if len(word) < 2:
                continue
                
            # 转小写比较
            word_lower = word.lower()
                
            # 停用词过滤
            if word_lower in stop_words:
                continue
                
            # 过滤"第X页"模式
            if re.match(r'^第[0-9一二三四五六七八九十百千]+页$', word):
                continue
                
            # 过滤纯数字（保留数字+单位的组合）
            if re.match(r'^[0-9]+$', word):
                continue
                
            # 过滤单字母或两字母的OCR噪音
            if re.match(r'^[a-zA-Z]{1,2}$', word) and word_lower not in ['id', 'ai', 'vr', 'ar', 'it', 'ui', 'ux']:
                continue
                
            # 过滤重复字符（如：aaa、111、中中中）
            if len(set(word)) == 1 and len(word) >= 3:
                continue
                
            filtered.append(word)
            
        return filtered
        
    def _enhance_with_terminology(self, words, original_text):
        """识别专业术语和行业关键词 - 提升其权重"""
        # 专业术语词典（可根据业务扩展）
        professional_terms = {
            # 演出娱乐行业
            '演出', '表演', '节目', '艺人', '艺能', '经纪', '演员', '歌手', '舞蹈',
            '魔术', '杂技', '话剧', '音乐会', '演唱会', '舞台', '灯光', '音响',
            '导演', '制作', '策划', '编排', '排练', '彩排',
                
            # 证件证书类
            '身份证', '护照', '签证', '许可证', '资格证', '证书', '执照',
            '演出经纪人证', '经纪人证书', '从业资格', '职业资格',
                
            # 合同财务类
            '合同', '协议', '同意函', '申请表', '付款', '发票', '报销',
            '费用', '金额', '结算', '账单', '税费', '预算',
                
            # 行政审批类
            '批文', '审批', '报批', '备案', '登记', '申请', '许可',
            '通知', '公示', '公告', '函件', '证明',
                
            # 场馆相关
            '长隆', '海洋王国', '主题公园', '景区', '场馆', '剧场', '舞台',
                
            # 时间相关（重要）
            '年度', '季度', '月度', '每日', '日期', '时间', '期限', '有效期',
        }
            
        enhanced = {}
        for word in words:
            # 检查是否为专业术语
            is_professional = False
            for term in professional_terms:
                if term in word or word in term:
                    is_professional = True
                    break
                
            # 专业术语权重加倍
            if is_professional:
                enhanced[word] = enhanced.get(word, 0) + 2
            else:
                enhanced[word] = enhanced.get(word, 0) + 1
            
        return enhanced
        
    def _calculate_word_scores(self, word_weights, original_text):
        """计算词汇权重分数 - 综合考虑多个因素"""
        scores = {}
        text_lower = original_text.lower()
        text_length = len(original_text)
            
        for word, base_weight in word_weights.items():
            score = base_weight
                
            # 1. 词频加权（出现越多越重要）
            word_count = text_lower.count(word.lower())
            score += word_count * 0.5
                
            # 2. 位置加权（出现在前面的更重要）
            first_position = text_lower.find(word.lower())
            if first_position >= 0:
                position_weight = 1.0 - (first_position / text_length)
                score += position_weight * 0.3
                
            # 3. 长度加权（适当的长度更有意义）
            word_len = len(word)
            if 2 <= word_len <= 4:  # 中文2-4字词最常见
                score += 0.5
            elif 4 < word_len <= 6:
                score += 0.3
            elif word_len > 8:  # 过长可能是错误识别
                score -= 0.2
                
            # 4. 中英文混合词（通常是专业术语）
            if re.search(r'[\u4e00-\u9fff]', word) and re.search(r'[A-Za-z]', word):
                score += 1.0
                
            # 5. 全大写英文（通常是缩写或重要标识）
            if word.isupper() and len(word) >= 2:
                score += 0.8
                
            scores[word] = score
            
        return scores
    
    def _merge_keywords(self, metadata_keywords, content_keywords, max_total=30):
        """按优先级合并关键词：
        1) 先使用文件名/路径/语义标签得到的元数据关键词
        2) 不足时再补充正文内容关键词
        """
        final = []
        seen = set()
        
        # 1. 元数据关键词优先
        for kw in metadata_keywords or []:
            if not kw:
                continue
            key = str(kw).lower()
            if key in seen:
                continue
            final.append(kw)
            seen.add(key)
            if len(final) >= max_total:
                return final
        
        # 2. 正文关键词作为补充
        for kw in content_keywords or []:
            if not kw:
                continue
            key = str(kw).lower()
            if key in seen:
                continue
            final.append(kw)
            seen.add(key)
            if len(final) >= max_total:
                break
        
        return final
    
    def extract_file_metadata(self, file_path):
        """提取文件元数据信息 - 增强版（为 AI 智能体提供更多信息）"""
        metadata = {
            'file_name': '',
            'file_dir': '',
            'file_ext': '',
            'file_size': 0,
            'file_size_formatted': '',
            'parent_dirs': [],
            'full_path': '',
            # 图片特有属性
            'image_width': None,
            'image_height': None,
            'image_format': None,
            'image_mode': None,
            # 时间信息
            'created_time': None,
            'modified_time': None,
            # 结构化标签信息（为 AI 智能体使用）
            'tags': [],
            'semantic_labels': {}
        }
        
        try:
            # 基本信息
            metadata['full_path'] = file_path
            metadata['file_name'] = os.path.basename(file_path)
            metadata['file_dir'] = os.path.dirname(file_path)
            metadata['file_ext'] = Path(file_path).suffix.lower()
            metadata['file_size'] = os.path.getsize(file_path)
            metadata['file_size_formatted'] = self._format_file_size(metadata['file_size'])
            # 创建/修改时间
            try:
                stat_res = os.stat(file_path)
                metadata['created_time'] = datetime.fromtimestamp(stat_res.st_ctime).isoformat()
                metadata['modified_time'] = datetime.fromtimestamp(stat_res.st_mtime).isoformat()
            except Exception:
                metadata['created_time'] = None
                metadata['modified_time'] = None
            
            # 提取父级目录名称（最多3级）
            path_parts = Path(file_path).parts
            # 过滤掉根目录和文件名
            dir_parts = [p for p in path_parts[:-1] if p != '/' and p != '']
            # 取最后3级目录
            metadata['parent_dirs'] = dir_parts[-3:] if len(dir_parts) >= 3 else dir_parts
            
            # 如果是图片文件，提取图片属性
            if metadata['file_ext'] in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif', '.webp']:
                try:
                    from PIL import Image
                    with Image.open(file_path) as img:
                        metadata['image_width'] = img.width
                        metadata['image_height'] = img.height
                        metadata['image_format'] = img.format
                        metadata['image_mode'] = img.mode  # RGB, RGBA, L 等
                except Exception as e:
                    print(f"提取图片属性失败: {e}")
            
            # 提取语义标签
            semantic_info = self._extract_semantic_info(metadata)
            metadata['semantic_labels'] = semantic_info
            
            # 从文件名和目录提取标签
            tags = []
            # 文件名标签
            file_name_without_ext = os.path.splitext(metadata['file_name'])[0]
            name_parts = re.split(r'[\s_\-。，；：“”‘’～！？、《》〈〉『』「」【】…—（）]+', file_name_without_ext)
            tags.extend([p.strip() for p in name_parts if p.strip() and len(p.strip()) > 1])
            
            # 目录标签
            tags.extend(metadata['parent_dirs'])
            
            # 去重
            metadata['tags'] = list(dict.fromkeys([tag for tag in tags if tag]))
            
        except Exception as e:
            print(f"提取文件元数据失败: {e}")
        
        return metadata
    
    def _build_metadata_content(self, metadata):
        """根据元数据构建内容 - 增强语义识别（为AI智能体优化）"""
        content_parts = []
        
        # 智能提取语义信息
        semantic_info = metadata.get('semantic_labels', {})
        
        # === 基础信息 ===
        content_parts.append("=== 文件基础信息 ===")
        
        # 文件名（去掉扩展名）
        file_name = metadata.get('file_name', '')
        if file_name:
            name_without_ext = os.path.splitext(file_name)[0]
            content_parts.append(f"文件名: {name_without_ext}")
        
        # 文件类型
        file_ext = metadata.get('file_ext', '')
        if file_ext:
            content_parts.append(f"文件格式: {file_ext}")
        
        # 文件大小
        file_size_formatted = metadata.get('file_size_formatted', '')
        if file_size_formatted:
            content_parts.append(f"文件大小: {file_size_formatted}")
        
        # 创建/修改时间
        created_time = metadata.get('created_time')
        modified_time = metadata.get('modified_time')
        if created_time:
            content_parts.append(f"创建时间: {created_time}")
        if modified_time:
            content_parts.append(f"修改时间: {modified_time}")
        
        # === 图片属性 ===
        if metadata.get('image_width') and metadata.get('image_height'):
            content_parts.append("\n=== 图片属性 ===")
            content_parts.append(f"图片尺寸: {metadata['image_width']} x {metadata['image_height']} 像素")
            
            # 计算分辨率类别
            width = metadata['image_width']
            height = metadata['image_height']
            if width >= 1920 and height >= 1080:
                resolution = "Full HD 或更高"
            elif width >= 1280 and height >= 720:
                resolution = "HD"
            else:
                resolution = "标准分辨率"
            content_parts.append(f"分辨率类别: {resolution}")
            
            if metadata.get('image_format'):
                content_parts.append(f"图片格式: {metadata['image_format']}")
            
            if metadata.get('image_mode'):
                mode_desc = {
                    'RGB': '彩色',
                    'RGBA': '彩色(透明)',
                    'L': '灰度',
                    'CMYK': 'CMYK印刷色彩'
                }.get(metadata['image_mode'], metadata['image_mode'])
                content_parts.append(f"色彩模式: {mode_desc}")
        
        # === 语义信息 ===
        has_semantic = any(semantic_info.values())
        if has_semantic:
            content_parts.append("\n=== 智能识别信息 ===")
            
            # 人员姓名
            if semantic_info.get('person_name'):
                content_parts.append(f"人员姓名: {semantic_info['person_name']}")
            
            # 文件类型/证书类型
            if semantic_info.get('document_type'):
                content_parts.append(f"文档类型: {semantic_info['document_type']}")
            
            # 机构名称
            if semantic_info.get('organization'):
                content_parts.append(f"相关机构: {semantic_info['organization']}")
            
            # 目录分类
            if semantic_info.get('category'):
                content_parts.append(f"目录分类: {semantic_info['category']}")
        
        # === 路径信息 ===
        content_parts.append("\n=== 路径信息 ===")
        
        # 所在目录
        parent_dirs = metadata.get('parent_dirs', [])
        if parent_dirs:
            dir_path = ' > '.join(parent_dirs)
            content_parts.append(f"所在目录: {dir_path}")
        
        # 完整路径
        full_path = metadata.get('full_path', '')
        if full_path:
            content_parts.append(f"完整路径: {full_path}")
        
        # === 标签信息 (为AI智能体使用) ===
        tags = metadata.get('tags', [])
        if tags:
            content_parts.append("\n=== 关键标签 ===")
            content_parts.append(f"标签: {', '.join(tags[:15])}")
        
        return '\n'.join(content_parts) if content_parts else '文件信息: 无法提取'
    
    def get_structured_metadata_for_ai(self, file_path):
        """
        为AI智能体生成结构化的元数据 (JSON格式)
        当OCR识别失败时，可以将此数据传递给Dify AI进行分析
        """
        import json
        
        try:
            # 提取完整元数据
            metadata = self.extract_file_metadata(file_path)
            
            # 构建为AI优化的结构
            ai_metadata = {
                # 基础信息
                "file_info": {
                    "name": os.path.splitext(metadata.get('file_name', ''))[0],
                    "extension": metadata.get('file_ext', ''),
                    "size_bytes": metadata.get('file_size', 0),
                    "size_formatted": metadata.get('file_size_formatted', ''),
                    "full_path": metadata.get('full_path', '')
                },
                
                # 图片属性 (如果有)
                "image_properties": None,
                
                # 语义标签
                "semantic_labels": metadata.get('semantic_labels', {}),
                
                # 路径信息
                "path_info": {
                    "directory": metadata.get('file_dir', ''),
                    "parent_directories": metadata.get('parent_dirs', []),
                    "directory_path": ' > '.join(metadata.get('parent_dirs', []))
                },
                
                # 关键标签
                "tags": metadata.get('tags', []),
                
                # OCR状态
                "ocr_status": "failed",  # 调用方可以更新
                
                # 上http://
                "context": {
                    "description": f"此文件位于 '{' > '.join(metadata.get('parent_dirs', []))}' 目录",
                    "file_type_hint": self._get_file_type_hint(metadata),
                    "possible_content": self._guess_content_from_metadata(metadata)
                }
            }
            
            # 如果是图片，添加图片属性
            if metadata.get('image_width'):
                ai_metadata["image_properties"] = {
                    "width": metadata['image_width'],
                    "height": metadata['image_height'],
                    "format": metadata.get('image_format', ''),
                    "mode": metadata.get('image_mode', ''),
                    "resolution_category": self._get_resolution_category(metadata['image_width'], metadata['image_height']),
                    "aspect_ratio": f"{metadata['image_width']}:{metadata['image_height']}"
                }
            
            return ai_metadata
            
        except Exception as e:
            print(f"生成AI元数据失败: {e}")
            return None
    
    def _get_resolution_category(self, width, height):
        """获取分辨率类别"""
        if width >= 3840 and height >= 2160:
            return "4K Ultra HD"
        elif width >= 1920 and height >= 1080:
            return "Full HD (1080p)"
        elif width >= 1280 and height >= 720:
            return "HD (720p)"
        else:
            return "Standard Definition"
    
    def _get_file_type_hint(self, metadata):
        """根据元数据推测文件类型提示"""
        semantic = metadata.get('semantic_labels', {})
        file_name = metadata.get('file_name', '')
        
        if semantic.get('document_type'):
            return f"{semantic['document_type']}类文件"
        elif '封面' in file_name:
            return "封面图片或海报"
        elif '证件' in file_name or '证书' in file_name:
            return "证件或证书类文件"
        elif metadata.get('image_width'):
            return "图片文件"
        else:
            return "文档文件"
    
    def _guess_content_from_metadata(self, metadata):
        """根据元数据推测可能的内容"""
        semantic = metadata.get('semantic_labels', {})
        suggestions = []
        
        if semantic.get('person_name'):
            suggestions.append(f"可能包含{semantic['person_name']}的信息")
        
        if semantic.get('organization'):
            suggestions.append(f"与{semantic['organization']}相关")
        
        if semantic.get('category'):
            suggestions.append(f"属于{semantic['category']}分类")
        
        file_name = metadata.get('file_name', '')
        if '封面' in file_name:
            suggestions.append("可能是作品或活动的封面图")
        elif '证件' in file_name:
            suggestions.append("可能包含身份证件或资质证书")
        
        return suggestions if suggestions else ["需要AI进一步分析图片内容"]
    
    def _extract_metadata_keywords(self, metadata):
        """从元数据提取关键词 - 增强语义识别"""
        keywords = []
        
        # 提取语义信息
        semantic_info = self._extract_semantic_info(metadata)
        
        # 优先添加语义关键词
        if semantic_info.get('person_name'):
            keywords.append(semantic_info['person_name'])
        
        if semantic_info.get('organization'):
            keywords.append(semantic_info['organization'])
        
        if semantic_info.get('document_type'):
            keywords.append(semantic_info['document_type'])
        
        if semantic_info.get('category'):
            keywords.append(semantic_info['category'])
        
        # 从文件名提取（去掉扩展名）
        file_name = metadata.get('file_name', '')
        if file_name:
            name_without_ext = os.path.splitext(file_name)[0]
            # 分割文件名（按空格、下划线、中文符号等）
            name_parts = re.split(r'[\s_\-。，；：“”‘’～！？、《》〈〉『』「」【】…—（）【】《》]+', name_without_ext)
            keywords.extend([p.strip() for p in name_parts if p.strip() and len(p.strip()) > 1])
        
        # 添加文件类型
        file_ext = metadata.get('file_ext', '')
        if file_ext:
            keywords.append(file_ext.replace('.', ''))
        
        # 从目录名提取
        parent_dirs = metadata.get('parent_dirs', [])
        for dir_name in parent_dirs:
            # 先添加完整目录名
            keywords.append(dir_name)
            # 再分割目录名
            dir_parts = re.split(r'[\s_\-]+', dir_name)
            keywords.extend([p.strip() for p in dir_parts if p.strip() and len(p.strip()) > 1])
        
        # 去重并限制数量
        unique_keywords = []
        seen = set()
        for kw in keywords:
            kw_lower = kw.lower()
            if kw_lower not in seen and len(kw) > 1:
                unique_keywords.append(kw)
                seen.add(kw_lower)
                if len(unique_keywords) >= 30:  # 增加到30个
                    break
        
        return unique_keywords
    
    def _extract_semantic_info(self, metadata):
        """从元数据中提取语义信息 - 增强版，支持更精细的分类"""
        semantic_info = {
            'person_name': None,         # 人员姓名
            'organization': None,        # 机构名称
            'document_type': None,       # 文档类型
            'category': None,            # 目录分类
            'main_theme': None,          # 主题
            'sub_theme': [],             # 子主题
            'business_domain': None,     # 业务领域
            'content_tags': [],          # 内容标签
            'time_info': None,           # 时间信息
        }
        
        file_name = metadata.get('file_name', '')
        parent_dirs = metadata.get('parent_dirs', [])
        full_path = metadata.get('full_path', '')
        
        # 1. 提取业务领域（从路径和文件名）
        business_domains = {
            '演出': ['演出', '表演', '节目', '演员', '艺人', '舞台'],
            '财务': ['财务', '付款', '报销', '发票', '费用', '金额', '账单'],
            '人事': ['人事', '招聘', '合同', '员工', '考勤', '薪资'],
            '法务': ['法务', '合同', '协议', '法律', '诉讼', '纠纷'],
            '行政': ['行政', '审批', '报批', '备案', '登记', '批文'],
            '营销': ['营销', '宣传', '推广', '活动', '海报', '广告'],
            '运营': ['运营', '管理', '调度', '统计', '报表', '数据'],
        }
        
        for domain, keywords in business_domains.items():
            for keyword in keywords:
                if keyword in full_path or keyword in file_name:
                    semantic_info['business_domain'] = domain
                    break
            if semantic_info['business_domain']:
                break
        
        # 2. 提取机构名称（扩展）
        organization_keywords = [
            '长隆海洋王国',
            '长隆',
            '海洋王国',
            '珠海长隆',
            '艺能娱乐',
            '艺能制作',
            'Impact',
        ]
        
        for keyword in organization_keywords:
            if keyword in full_path or keyword in file_name:
                semantic_info['organization'] = keyword
                break
        
        # 3. 提取目录分类（增强版）
        category_mapping = {
            # 一级分类 -> 二级分类
            '艺人资料': ['艺人证件', '艺人照片', '艺人简历', '艺人合同'],
            '演出资料': ['节目策划', '演出方案', '排练记录', '演出合同'],
            '证件资料': ['身份证', '护照', '证书', '许可证'],
            '合同协议': ['演出合同', '服务协议', '采购合同', '租赁合同'],
            '财务资料': ['付款申请', '报销单据', '发票', '费用清单'],
            '审批文件': ['批文', '审批表', '报批资料', '备案证明'],
        }
        
        for main_cat, sub_cats in category_mapping.items():
            # 检查主分类
            for dir_name in parent_dirs:
                if main_cat in dir_name:
                    semantic_info['category'] = main_cat
                    # 检查子分类
                    for sub_cat in sub_cats:
                        if sub_cat in dir_name or sub_cat in file_name:
                            semantic_info['sub_theme'] = [sub_cat]
                            break
                    break
            if semantic_info['category']:
                break
        
        # 如果没有找到主分类，尝试直接从文件名匹配
        if not semantic_info['category']:
            for main_cat, sub_cats in category_mapping.items():
                for sub_cat in sub_cats:
                    if sub_cat in file_name:
                        semantic_info['category'] = main_cat
                        semantic_info['sub_theme'] = [sub_cat]
                        break
                if semantic_info['category']:
                    break
        
        # 4. 提取文档类型（增强版）
        document_types = {
            # 证件类
            '身份证': ['身份证', 'ID', 'idcard'],
            '护照': ['护照', 'passport'],
            '签证': ['签证', 'visa'],
            '演出经纪人证书': ['演出经纪人证书', '演出经纪人证', '经纪人证书', '经纪人证'],
            '资格证书': ['资格证', '证书', '职业证'],
            
            # 合同类
            '演出合同': ['演出合同', '表演合同'],
            '服务合同': ['服务合同', '服务协议'],
            '劳动合同': ['劳动合同', '雇佣合同'],
            
            # 文书类
            '同意函': ['同意函', '演出同意函'],
            '付款申请': ['付款申请', '付款申请表'],
            '批文': ['批文', '审批文件'],
            '通知': ['通知', '通知函'],
            '证明': ['证明', '证明文件'],
        }
        
        for doc_type, patterns in document_types.items():
            for pattern in patterns:
                if pattern.lower() in file_name.lower():
                    semantic_info['document_type'] = doc_type
                    break
            if semantic_info['document_type']:
                break
        
        # 5. 提取主题信息
        # 主题通常是文件名中最显著的部分
        name_without_ext = os.path.splitext(file_name)[0]
        # 移除日期、编号等信息后的主体内容
        theme = re.sub(r'[0-9]{4}[-_]?[0-9]{2}[-_]?[0-9]{2}', '', name_without_ext)  # 移除日期
        theme = re.sub(r'[0-9]{6,}', '', theme)  # 移除长数字编号
        theme = re.sub(r'[【】\[\]\(\)（）]', ' ', theme)  # 移除括号
        theme = theme.strip()
        if len(theme) > 2:
            semantic_info['main_theme'] = theme
        
        # 6. 提取人员姓名（增强版）
        person_name = self._extract_person_name(file_name, parent_dirs)
        if person_name:
            semantic_info['person_name'] = person_name
        
        # 7. 提取时间信息
        time_info = self._extract_time_info(file_name, metadata)
        if time_info:
            semantic_info['time_info'] = time_info
        
        # 8. 生成内容标签（综合信息）
        content_tags = []
        if semantic_info['business_domain']:
            content_tags.append(semantic_info['business_domain'])
        if semantic_info['document_type']:
            content_tags.append(semantic_info['document_type'])
        if semantic_info['category']:
            content_tags.append(semantic_info['category'])
        if semantic_info['organization']:
            content_tags.append(semantic_info['organization'])
        semantic_info['content_tags'] = content_tags
        
        return semantic_info
    
    def _extract_person_name(self, file_name, parent_dirs):
        """从文件名或路径提取人员姓名"""
        # 常见中文姓氏（扩展）
        chinese_surnames = [
            '李', '王', '张', '刘', '陈', '杨', '黄', '赵', '吴', '周',
            '徐', '孙', '马', '朱', '胡', '郭', '何', '罗', '高', '林',
            '郑', '谢', '韩', '唐', '冯', '于', '董', '萧', '程', '曹',
            '袁', '邓', '许', '傅', '沈', '曾', '彭', '吕', '苏', '卢',
            '蒋', '蔡', '贾', '丁', '魏', '薛', '叶', '阎', '余', '潘',
            '杜', '戴', '夏', '钟', '汪', '田', '任', '姜', '范', '方',
            '石', '姚', '谭', '廖', '邹', '熊', '金', '陆', '郝', '孔',
            '白', '崔', '康', '毛', '邱', '秦', '江', '史', '顾', '侯',
            '邵', '孟', '龙', '万', '段', '汤', '黎', '易', '常',
            '谷', '关', '大', '小',
        ]
        
        # 排除常见非姓名词汇（扩展）
        exclude_words = [
            '申请', '合同', '协议', '证书', '通知', '批文',
            '长隆', '海洋', '王国', '海洋王国', '长隆海洋',
            '艺能', '娱乐', '演出', '表演', '节目',
            '成龙', '李小龙',  # 明星名字不计入（需要特殊处理）
        ]
        
        # 尝试匹配中文姓名模式（姓氏 + 1-3个字）
        name_pattern = r'([' + ''.join(chinese_surnames) + r'][\u4e00-\u9fff]{1,3})'
        
        # 从文件名中提取
        name_matches = re.findall(name_pattern, file_name)
        
        # 过滤掉非姓名词汇
        valid_names = []
        for name in name_matches:
            is_valid = True
            # 检查是否在排除列表中
            for exclude in exclude_words:
                if name in exclude:
                    is_valid = False
                    break
            # 检查长度（2-4个字为佳）
            if len(name) < 2 or len(name) > 4:
                is_valid = False
            if is_valid:
                valid_names.append(name)
        
        if valid_names:
            # 选择最可能的姓名（通常是最后一个）
            return valid_names[-1]
        
        # 如果文件名没有，尝试从目录名提取
        for dir_name in parent_dirs:
            name_matches = re.findall(name_pattern, dir_name)
            valid_names = [name for name in name_matches 
                          if not any(name in exclude for exclude in exclude_words)
                          and 2 <= len(name) <= 4]
            if valid_names:
                return valid_names[-1]
        
        return None
    
    def _extract_time_info(self, file_name, metadata):
        """提取时间信息"""
        time_info = {}
        
        # 1. 从文件名提取日期（2023-01-01, 20230101, 2023年1月1日）
        date_patterns = [
            (r'(\d{4})[-_](\d{1,2})[-_](\d{1,2})', 'date'),  # 2023-01-01
            (r'(\d{4})(\d{2})(\d{2})', 'date'),              # 20230101
            (r'(\d{4})年(\d{1,2})月(\d{1,2})日', 'date'),        # 2023年1月1日
            (r'(\d{4})年(\d{1,2})月', 'month'),                # 2023年1月
            (r'(\d{4})年', 'year'),                            # 2023年
        ]
        
        for pattern, time_type in date_patterns:
            match = re.search(pattern, file_name)
            if match:
                if time_type == 'date':
                    year, month, day = match.groups()
                    time_info['year'] = year
                    time_info['month'] = month.zfill(2)
                    time_info['day'] = day.zfill(2)
                    time_info['full_date'] = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                elif time_type == 'month':
                    year, month = match.groups()
                    time_info['year'] = year
                    time_info['month'] = month.zfill(2)
                elif time_type == 'year':
                    time_info['year'] = match.group(1)
                break
        
        # 2. 如果文件名没有日期，使用文件修改时间
        if not time_info and metadata.get('modified_time'):
            modified_time = metadata['modified_time']
            # modified_time 格式：2023-01-01T12:30:45
            date_match = re.match(r'(\d{4})-(\d{2})-(\d{2})', modified_time)
            if date_match:
                year, month, day = date_match.groups()
                time_info['year'] = year
                time_info['month'] = month
                time_info['day'] = day
                time_info['full_date'] = f"{year}-{month}-{day}"
                time_info['source'] = 'file_modified_time'
        
        return time_info if time_info else None
    
    def get_file_classification(self, file_path):
        """获取文件的精细分类信息 - 新增功能"""
        try:
            metadata = self.extract_file_metadata(file_path)
            semantic_info = metadata.get('semantic_labels', {})
            
            classification = {
                # 一级分类：业务领域
                'business_domain': semantic_info.get('business_domain', '未分类'),
                
                # 二级分类：文档类别
                'category': semantic_info.get('category', '未分类'),
                
                # 三级分类：文档类型
                'document_type': semantic_info.get('document_type', '未分类'),
                
                # 主题信息
                'main_theme': semantic_info.get('main_theme', ''),
                'sub_theme': semantic_info.get('sub_theme', []),
                
                # 关联信息
                'person': semantic_info.get('person_name', ''),
                'organization': semantic_info.get('organization', ''),
                'time_info': semantic_info.get('time_info', {}),
                
                # 标签
                'tags': semantic_info.get('content_tags', []),
                
                # 文件类型（技术类型）
                'file_type': self._get_technical_file_type(metadata),
            }
            
            return classification
            
        except Exception as e:
            print(f"获取文件分类失败: {e}")
            return None
    
    def _get_technical_file_type(self, metadata):
        """获取文件的技术类型"""
        ext = metadata.get('file_ext', '').lower()
        
        type_mapping = {
            # 图片类
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif', '.webp'],
            # 文档类
            'document': ['.pdf', '.doc', '.docx', '.txt', '.md', '.rtf'],
            # 表格类
            'spreadsheet': ['.xlsx', '.xls', '.csv'],
            # 演示文稿
            'presentation': ['.ppt', '.pptx'],
            # 代码类
            'code': ['.py', '.js', '.java', '.cpp', '.c', '.go', '.sh', '.html', '.css'],
            # 压缩文件
            'archive': ['.zip', '.rar', '.7z', '.tar', '.gz'],
            # 视频类
            'video': ['.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv'],
            # 音频类
            'audio': ['.mp3', '.wav', '.flac', '.aac', '.ogg'],
        }
        
        for file_type, extensions in type_mapping.items():
            if ext in extensions:
                return file_type
        
        return 'other'
    
    def _filter_watermark(self, content):
        """增强的水印和无意义内容过滤器"""
        if not content:
            return content
        
        # 第一步：过滤水印内容
        content = self._remove_watermarks(content)
        
        # 第二步：OCR错误纠正
        content = self._correct_ocr_errors(content)
        
        # 第三步：内容去重
        content = self._deduplicate_content(content)
        
        # 第四步：段落结构优化
        content = self._optimize_paragraph_structure(content)
        
        # 第五步：敏感信息脱敏（如需要）
        content = self._mask_sensitive_info(content)
        
        # 如果过滤后内容太少，返回空
        if len(content.strip()) < 20:
            return ''
        
        return content
    
    def _remove_watermarks(self, content):
        """移除水印内容 - 增强版"""
        if not content:
            return content
        
        # 扩展的水印关键词库
        watermark_patterns = [
            # 固定水印文本
            r'仅供.*?使用',
            r'报批使用',
            r'内部资料',
            r'机密文件',
            r'严禁外传',
            r'版权所有',
            r'copyright',
            r'confidential',
            
            # 重复的短语（通常是水印）
            r'(长隆海洋王国[\s]*){2,}',
            r'(海洋.*?烟花.*?){2,}',
            r'(Impact[\s]*){3,}',
            
            # 水印常见位置标记
            r'\[水印\]',
            r'【水印】',
            r'\(水印\)',
            r'（水印）',
        ]
        
        lines = content.split('\n')
        filtered_lines = []
        
        for line in lines:
            # 检查是否匹配水印模式
            is_watermark = False
            for pattern in watermark_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    is_watermark = True
                    break
            
            # 检查是否是重复行（水印特征）
            if not is_watermark:
                line_stripped = line.strip()
                # 如果同一行在文本中出现超过3次，可能是水印
                if line_stripped and content.count(line_stripped) > 3:
                    is_watermark = True
            
            # 检查是否是过短的行（少于5个字符，可能是噪音）
            if not is_watermark and line.strip():
                if len(line.strip()) < 5:
                    # 检查是否包含有效信息
                    if not re.search(r'[\u4e00-\u9fff]{2,}|[A-Z]{2,}|\d{2,}', line):
                        is_watermark = True
            
            # 如果不是水印，保留
            if not is_watermark and line.strip():
                filtered_lines.append(line)
        
        return '\n'.join(filtered_lines)
    
    def _correct_ocr_errors(self, content):
        """OCR错误纠正 - 修复常见识别错误"""
        if not content:
            return content
        
        # OCR常见错误映射（基于实际经验积累）
        error_corrections = {
            # 数字误识别
            'O': '0',  # 字母O -> 数字0（在数字上下文中）
            'l': '1',  # 小写L -> 数字1（在数字上下文中）
            'Z': '2',  # 在某些字体中
            'S': '5',  # 在某些字体中
            
            # 中文常见误识别
            '璋': '璨',  # 常见混淆
            '洊': '王',  # 常见混淆
            '干杯': '千杯',  # 语义纠正
            
            # 标点符号误识别
            '、、': '、',
            '。。': '。',
            ',,': ',',
            '..': '.',
            
            # 空格和换行符异常
            '  ': ' ',  # 多余空格
            '\n\n\n': '\n\n',  # 多余换行
        }
        
        corrected = content
        for error, correction in error_corrections.items():
            corrected = corrected.replace(error, correction)
        
        # 修复数字上下文中的字母
        # 例如：2O23 -> 2023, 1l5 -> 115
        corrected = re.sub(r'(\d)O(\d)', r'\g<1>0\g<2>', corrected)
        corrected = re.sub(r'(\d)l(\d)', r'\g<1>1\g<2>', corrected)
        
        # 移除OCR产生的重复字符
        # 例如：姓姓名名 -> 姓名
        corrected = re.sub(r'([\u4e00-\u9fff])\1{2,}', r'\1', corrected)
        
        return corrected
    
    def _deduplicate_content(self, content):
        """内容去重 - 移除重复段落和行"""
        if not content:
            return content
        
        lines = content.split('\n')
        unique_lines = []
        seen = set()
        
        for line in lines:
            line_stripped = line.strip()
            
            # 空行保留（但不计入去重）
            if not line_stripped:
                unique_lines.append('')
                continue
            
            # 太短的行（少于3个字符）总是保留
            if len(line_stripped) < 3:
                unique_lines.append(line)
                continue
            
            # 标准化后比较（忽略空格和标点差异）
            normalized = re.sub(r'[\s\W]+', '', line_stripped)
            
            # 如果未见过，添加
            if normalized not in seen:
                unique_lines.append(line)
                seen.add(normalized)
        
        return '\n'.join(unique_lines)
    
    def _optimize_paragraph_structure(self, content):
        """优化段落结构 - 保持内容逻辑性和可读性"""
        if not content:
            return content
        
        lines = content.split('\n')
        optimized_lines = []
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # 空行处理：连续空行压缩为单个空行
            if not line_stripped:
                if optimized_lines and optimized_lines[-1] != '':
                    optimized_lines.append('')
                continue
            
            # 检查是否应该合并到上一行
            # 规则：如果当前行很短（<10字符）且不是标题，可能应该合并
            if (optimized_lines and 
                len(line_stripped) < 10 and 
                not re.match(r'^[一二三四五六七八九十0-9]+[、.]', line_stripped) and
                not re.match(r'^[（(].*?[)）]', line_stripped)):
                
                last_line = optimized_lines[-1]
                # 如果上一行也很短，合并
                if last_line and len(last_line.strip()) < 50:
                    optimized_lines[-1] = last_line + ' ' + line_stripped
                    continue
            
            optimized_lines.append(line)
        
        # 移除首尾空行
        while optimized_lines and not optimized_lines[0].strip():
            optimized_lines.pop(0)
        while optimized_lines and not optimized_lines[-1].strip():
            optimized_lines.pop()
        
        return '\n'.join(optimized_lines)
    
    def _mask_sensitive_info(self, content):
        """敏感信息脱敏 - 自动识别和处理敏感数据"""
        if not content:
            return content
        
        masked = content
        
        # 1. 身份证号脱敏（保留前6位和后4位）
        # 例如：110101199001011234 -> 110101********1234
        masked = re.sub(
            r'(\d{6})\d{8}(\d{4})',
            r'\1********\2',
            masked
        )
        
        # 2. 手机号脱敏（保留前3位和后4位）
        # 例如：13812345678 -> 138****5678
        masked = re.sub(
            r'(1[3-9]\d)\d{4}(\d{4})',
            r'\1****\2',
            masked
        )
        
        # 3. 银行卡号脱敏（保留前6位和后4位）
        # 例如：6222021234567890123 -> 622202*********0123
        masked = re.sub(
            r'(\d{6})\d{9,13}(\d{4})',
            r'\1*********\2',
            masked
        )
        
        # 4. 邮箱脱敏（保留前2个字符和域名）
        # 例如：example@gmail.com -> ex****@gmail.com
        masked = re.sub(
            r'([a-zA-Z0-9]{1,2})[a-zA-Z0-9._-]*?(@[a-zA-Z0-9.-]+)',
            r'\1****\2',
            masked
        )
        
        # 注意：这个功能是可选的，可以通过参数控制是否启用
        # 在某些场景下（如内部系统），可能不需要脱敏
        
        return masked
    
    def extract_keywords_with_hierarchy(self, text, max_keywords=30):
        """提取多层级关键词 - 分为主要、次要、辅助关键词"""
        if not text or len(text.strip()) == 0:
            return {
                'primary': [],      # 主要关键词（专业术语、高频词）
                'secondary': [],    # 次要关键词（中频词）
                'auxiliary': []     # 辅助关键词（低频词、上下文词）
            }
        
        try:
            # 获取所有关键词及其权重
            words = self._intelligent_tokenize(text)
            filtered_words = self._filter_stopwords(words)
            
            if not filtered_words:
                return {'primary': [], 'secondary': [], 'auxiliary': []}
            
            enhanced_words = self._enhance_with_terminology(filtered_words, text)
            word_scores = self._calculate_word_scores(enhanced_words, text)
            
            # 按权重排序
            sorted_keywords = sorted(word_scores.items(), key=lambda x: x[1], reverse=True)
            
            # 分层
            total = min(len(sorted_keywords), max_keywords)
            primary_count = int(total * 0.4)      # 40%为主要关键词
            secondary_count = int(total * 0.35)   # 35%为次要关键词
            auxiliary_count = total - primary_count - secondary_count  # 剩余为辅助关键词
            
            result = {
                'primary': [word for word, score in sorted_keywords[:primary_count]],
                'secondary': [word for word, score in sorted_keywords[primary_count:primary_count+secondary_count]],
                'auxiliary': [word for word, score in sorted_keywords[primary_count+secondary_count:total]]
            }
            
            return result
            
        except Exception as e:
            print(f"提取多层级关键词失败: {e}")
            return {'primary': [], 'secondary': [], 'auxiliary': []}
    
    def get_content_summary(self, file_path):
        """获取文件内容摘要 - 为搜索和展示优化"""
        try:
            result = self.process_file(file_path)
            
            if not result['success']:
                return None
            
            # 提取多层级关键词
            hierarchical_keywords = self.extract_keywords_with_hierarchy(
                result['content'], 
                max_keywords=30
            )
            
            summary = {
                # 基本信息
                'file_name': result['metadata'].get('file_name', ''),
                'file_type': result['classification'].get('file_type', 'unknown'),
                
                # 分类信息
                'business_domain': result['classification'].get('business_domain', ''),
                'category': result['classification'].get('category', ''),
                'document_type': result['classification'].get('document_type', ''),
                
                # 关键信息
                'person': result['classification'].get('person', ''),
                'organization': result['classification'].get('organization', ''),
                'time_info': result['classification'].get('time_info', {}),
                
                # 关键词（分层）
                'keywords': hierarchical_keywords,
                
                # 内容预览
                'preview': result['preview'],
                
                # 质量评分
                'quality_score': result['quality_score'],
                
                # 标签
                'tags': result['classification'].get('tags', []),
            }
            
            return summary
            
        except Exception as e:
            print(f"获取内容摘要失败: {e}")
            return None
    
    def batch_process_files(self, file_paths, max_workers=4):
        """批量处理文件 - 提升性能"""
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        results = {}
        
        try:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # 提交所有任务
                future_to_path = {
                    executor.submit(self.process_file, path): path 
                    for path in file_paths
                }
                
                # 收集结果
                for future in as_completed(future_to_path):
                    path = future_to_path[future]
                    try:
                        result = future.result(timeout=30)  # 30秒超时
                        results[path] = result
                    except Exception as e:
                        print(f"处理文件 {path} 失败: {e}")
                        results[path] = {
                            'success': False,
                            'error': str(e)
                        }
        
        except Exception as e:
            print(f"批量处理失败: {e}")
        
        return results
    
    def get_processing_stats(self, file_path):
        """获取处理统计信息 - 用于调试和监控"""
        import time
        
        stats = {
            'file_path': file_path,
            'start_time': time.time(),
            'end_time': None,
            'duration': None,
            'success': False,
            'content_length': 0,
            'keyword_count': 0,
            'quality_score': 0.0,
            'processing_method': '',  # 'text', 'ocr', 'metadata_only'
            'errors': []
        }
        
        try:
            result = self.process_file(file_path)
            
            stats['end_time'] = time.time()
            stats['duration'] = stats['end_time'] - stats['start_time']
            stats['success'] = result['success']
            stats['content_length'] = len(result.get('content', ''))
            stats['keyword_count'] = len(result.get('keywords', []))
            stats['quality_score'] = result.get('quality_score', 0.0)
            
            # 判断处理方法
            content = result.get('content', '')
            if '【内容提取失败】' in content:
                stats['processing_method'] = 'metadata_only'
            elif Path(file_path).suffix.lower() in ['.jpg', '.jpeg', '.png', '.pdf']:
                stats['processing_method'] = 'ocr'
            else:
                stats['processing_method'] = 'text'
            
            if 'error' in result:
                stats['errors'].append(result['error'])
            
        except Exception as e:
            stats['end_time'] = time.time()
            stats['duration'] = stats['end_time'] - stats['start_time']
            stats['errors'].append(str(e))
        
        return stats

# 创建全局实例
file_processor = FileProcessor()
