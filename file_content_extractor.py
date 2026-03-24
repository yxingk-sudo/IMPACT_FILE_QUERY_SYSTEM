#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件内容提取模块
支持多种文件格式的内容提取
"""

import os
import mimetypes
from typing import Dict, Optional, Tuple

class FileContentExtractor:
    """文件内容提取器"""
    
    def __init__(self):
        """初始化提取器"""
        self.supported_extensions = {
            '.txt': self._extract_text,
            '.md': self._extract_text,
            '.csv': self._extract_text,
            '.doc': self._extract_doc,
            '.docx': self._extract_docx,
            '.pdf': self._extract_pdf,
            '.jpg': self._extract_image,
            '.jpeg': self._extract_image,
            '.png': self._extract_image,
            '.bmp': self._extract_image,
            '.gif': self._extract_image
        }
    
    def extract_content(self, file_path: str) -> Tuple[bool, Optional[str], Optional[str]]:
        """
        提取文件内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            Tuple[bool, Optional[str], Optional[str]]: (成功标志, 内容, 错误信息)
        """
        try:
            if not os.path.exists(file_path):
                return False, None, f"文件不存在: {file_path}"
            
            if not os.path.isfile(file_path):
                return False, None, f"路径不是文件: {file_path}"
            
            # 获取文件扩展名
            _, ext = os.path.splitext(file_path.lower())
            
            # 检查是否支持该文件类型
            if ext not in self.supported_extensions:
                return False, None, f"不支持的文件类型: {ext}"
            
            # 调用对应的提取方法
            extractor = self.supported_extensions[ext]
            content = extractor(file_path)
            
            return True, content, None
            
        except Exception as e:
            return False, None, f"提取文件内容失败: {str(e)}"
    
    def _extract_text(self, file_path: str) -> str:
        """提取文本文件内容"""
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，返回错误信息
            return "[无法读取文本文件: 编码错误]"
            
        except Exception as e:
            return f"[文本文件读取错误: {str(e)}]"
    
    def _extract_doc(self, file_path: str) -> str:
        """提取.doc文件内容"""
        try:
            # 尝试使用python-docx处理.doc文件
            # 注意：python-docx只支持.docx文件，.doc文件需要其他库
            try:
                from docx import Document
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n".join(paragraphs)
            except Exception:
                # 如果失败，尝试使用其他方法或返回占位信息
                return f"[.DOC文件 - 文件名: {os.path.basename(file_path)}]"
                
        except Exception as e:
            return f"[DOC文件读取错误: {str(e)}]"
    
    def _extract_docx(self, file_path: str) -> str:
        """提取.docx文件内容"""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
            
        except Exception as e:
            return f"[DOCX文件读取错误: {str(e)}]"
    
    def _extract_pdf(self, file_path: str) -> str:
        """提取PDF文件内容"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text("text")
                if page_text.strip():
                    text.append(page_text)
            
            doc.close()
            return "\n".join(text)
            
        except Exception as e:
            return f"[PDF文件读取错误: {str(e)}]"
    
    def _extract_image(self, file_path: str) -> str:
        """提取图像文件内容（OCR）"""
        try:
            import pytesseract
            from PIL import Image
            
            # 打开图像
            img = Image.open(file_path)
            
            # 图像预处理
            img_gray = img.convert('L')
            
            # 使用Tesseract进行OCR
            text = pytesseract.image_to_string(img_gray, lang='chi_sim+chi_tra+eng')
            
            if text.strip():
                return text.strip()
            else:
                return f"[图像文件 - 文件名: {os.path.basename(file_path)}]"
                
        except Exception as e:
            return f"[图像文件读取错误: {str(e)}]"

# 创建全局实例
file_extractor = FileContentExtractor()
